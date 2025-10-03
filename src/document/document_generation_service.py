"""
Document Generation Service
Generates documents from AI responses and converts to various formats (PDF, DOCX, LaTeX)
"""

import os
import tempfile
import subprocess
import logging
from typing import Optional, Dict, Any, List
from pathlib import Path
from datetime import datetime
import re

# Document generation libraries
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("reportlab not available - PDF generation will be limited")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx not available - DOCX generation will be disabled")

logger = logging.getLogger(__name__)


class DocumentGenerationService:
    """Service for generating documents from AI responses"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.temp_dir = tempfile.gettempdir()
    
    def generate_latex_document(self, content: str, 
                                assignment_name: str = "Assignment",
                                author: str = "Student") -> str:
        """
        Generate LaTeX document from content
        
        Args:
            content: Content with LaTeX formatting
            assignment_name: Title of the assignment
            author: Author name
            
        Returns:
            LaTeX document as string
        """
        # Clean content and ensure proper LaTeX formatting
        latex_content = self._prepare_latex_content(content)
        
        latex_doc = f"""\\documentclass[12pt, a4paper]{{article}}

% Packages
\\usepackage{{amsmath}}
\\usepackage{{amssymb}}
\\usepackage{{graphicx}}
\\usepackage{{hyperref}}
\\usepackage{{geometry}}
\\usepackage{{fancyhdr}}
\\usepackage[utf8]{{inputenc}}
\\usepackage{{listings}}
\\usepackage{{xcolor}}

% Page layout
\\geometry{{margin=1in}}
\\pagestyle{{fancy}}
\\fancyhf{{}}
\\rhead{{\\thepage}}
\\lhead{{{assignment_name}}}

% Hyperref setup
\\hypersetup{{
    colorlinks=true,
    linkcolor=blue,
    urlcolor=blue,
    citecolor=blue
}}

% Title information
\\title{{{assignment_name}}}
\\author{{{author}}}
\\date{{\\today}}

\\begin{{document}}

\\maketitle
\\tableofcontents
\\newpage

{latex_content}

\\end{{document}}
"""
        return latex_doc
    
    def latex_to_pdf(self, latex_content: str, output_filename: str = None) -> Optional[str]:
        """
        Convert LaTeX to PDF using pdflatex
        
        Args:
            latex_content: LaTeX document content
            output_filename: Desired output filename (without extension)
            
        Returns:
            Path to generated PDF file or None if failed
        """
        if output_filename is None:
            output_filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        try:
            # Create temporary directory for LaTeX compilation
            with tempfile.TemporaryDirectory() as temp_dir:
                # Write LaTeX file
                tex_file = os.path.join(temp_dir, "document.tex")
                with open(tex_file, 'w', encoding='utf-8') as f:
                    f.write(latex_content)
                
                # Compile with pdflatex (run twice for TOC and references)
                for i in range(2):
                    result = subprocess.run(
                        ['pdflatex', '-interaction=nonstopmode', '-output-directory', temp_dir, tex_file],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    
                    if result.returncode != 0 and i == 1:
                        self.logger.error(f"pdflatex failed: {result.stderr}")
                        return None
                
                # Copy PDF to output location
                pdf_file = os.path.join(temp_dir, "document.pdf")
                if os.path.exists(pdf_file):
                    output_path = os.path.join(self.temp_dir, f"{output_filename}.pdf")
                    import shutil
                    shutil.copy2(pdf_file, output_path)
                    self.logger.info(f"PDF generated successfully: {output_path}")
                    return output_path
                else:
                    self.logger.error("PDF file not generated")
                    return None
                    
        except FileNotFoundError:
            self.logger.error("pdflatex not found. Please install TeX Live or MiKTeX")
            return None
        except subprocess.TimeoutExpired:
            self.logger.error("LaTeX compilation timed out")
            return None
        except Exception as e:
            self.logger.error(f"Error converting LaTeX to PDF: {e}")
            return None
    
    def generate_pdf_from_markdown(self, content: str, 
                                   assignment_name: str = "Assignment",
                                   author: str = "Student") -> Optional[str]:
        """
        Generate PDF from Markdown content using reportlab
        
        Args:
            content: Markdown formatted content
            assignment_name: Title of the assignment
            author: Author name
            
        Returns:
            Path to generated PDF file or None if failed
        """
        if not REPORTLAB_AVAILABLE:
            self.logger.error("reportlab not available")
            return None
        
        try:
            # Create output filename
            output_filename = f"{assignment_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            output_path = os.path.join(self.temp_dir, output_filename)
            
            # Create PDF document
            doc = SimpleDocTemplate(output_path, pagesize=letter,
                                  topMargin=1*inch, bottomMargin=1*inch)
            
            # Container for document elements
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=30,
                alignment=1  # Center
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor=colors.HexColor('#333333'),
                spaceAfter=12,
                spaceBefore=12
            )
            
            # Add title
            story.append(Paragraph(assignment_name, title_style))
            story.append(Paragraph(f"Author: {author}", styles['Normal']))
            story.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
            story.append(Spacer(1, 0.3*inch))
            
            # Parse and add content
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Handle headers
                    if para.startswith('##'):
                        text = para.lstrip('#').strip()
                        story.append(Paragraph(text, heading_style))
                    elif para.startswith('#'):
                        text = para.lstrip('#').strip()
                        story.append(Paragraph(text, title_style))
                    else:
                        # Handle bold and italic
                        para = self._convert_markdown_to_html(para)
                        story.append(Paragraph(para, styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            # Build PDF
            doc.build(story)
            self.logger.info(f"PDF generated successfully: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Error generating PDF from markdown: {e}")
            return None
    
    def generate_docx(self, content: str,
                     assignment_name: str = "Assignment",
                     author: str = "Student") -> Optional[str]:
        """
        Generate DOCX document from Markdown content
        
        Args:
            content: Markdown formatted content
            assignment_name: Title of the assignment
            author: Author name
            
        Returns:
            Path to generated DOCX file or None if failed
        """
        if not PYTHON_DOCX_AVAILABLE:
            self.logger.error("python-docx not available")
            return None
        
        try:
            # Create output filename
            output_filename = f"{assignment_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = os.path.join(self.temp_dir, output_filename)
            
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading(assignment_name, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            doc.add_paragraph(f"Author: {author}")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph()
            
            # Parse and add content
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Handle headers
                    if para.startswith('###'):
                        text = para.lstrip('#').strip()
                        doc.add_heading(text, 3)
                    elif para.startswith('##'):
                        text = para.lstrip('#').strip()
                        doc.add_heading(text, 2)
                    elif para.startswith('#'):
                        text = para.lstrip('#').strip()
                        doc.add_heading(text, 1)
                    else:
                        # Add paragraph with basic formatting
                        p = doc.add_paragraph()
                        self._add_formatted_text(p, para)
            
            # Save document
            doc.save(output_path)
            self.logger.info(f"DOCX generated successfully: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Error generating DOCX: {e}")
            return None
    
    def _prepare_latex_content(self, content: str) -> str:
        """Prepare content for LaTeX document"""
        # Convert markdown headers to LaTeX sections
        content = re.sub(r'^### (.+)$', r'\\subsubsection{\1}', content, flags=re.MULTILINE)
        content = re.sub(r'^## (.+)$', r'\\subsection{\1}', content, flags=re.MULTILINE)
        content = re.sub(r'^# (.+)$', r'\\section{\1}', content, flags=re.MULTILINE)
        
        # Convert markdown bold/italic to LaTeX
        content = re.sub(r'\*\*(.+?)\*\*', r'\\textbf{\1}', content)
        content = re.sub(r'\*(.+?)\*', r'\\textit{\1}', content)
        
        # Convert inline code to LaTeX
        content = re.sub(r'`(.+?)`', r'\\texttt{\1}', content)
        
        # Convert bullet points to LaTeX itemize
        content = re.sub(r'^- (.+)$', r'\\item \1', content, flags=re.MULTILINE)
        
        # Wrap itemize blocks
        lines = content.split('\n')
        in_list = False
        result = []
        for line in lines:
            if line.strip().startswith('\\item') and not in_list:
                result.append('\\begin{itemize}')
                in_list = True
            elif not line.strip().startswith('\\item') and in_list:
                result.append('\\end{itemize}')
                in_list = False
            result.append(line)
        if in_list:
            result.append('\\end{itemize}')
        
        return '\n'.join(result)
    
    def _convert_markdown_to_html(self, text: str) -> str:
        """Convert basic Markdown to HTML for reportlab"""
        # Bold
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        # Italic
        text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
        # Code
        text = re.sub(r'`(.+?)`', r'<font face="Courier">\1</font>', text)
        return text
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add formatted text to DOCX paragraph"""
        # Simple formatting - just add as plain text for now
        # More sophisticated parsing could be added
        paragraph.add_run(text)


# Example usage
if __name__ == "__main__":
    service = DocumentGenerationService()
    
    # Test content
    content = """
# Introduction

This is a **sample** assignment with *formatting*.

## Methods

- Point 1
- Point 2
- Point 3

## Results

The equation is $E = mc^2$.
"""
    
    # Generate LaTeX
    latex_doc = service.generate_latex_document(content, "Sample Assignment", "John Doe")
    print("LaTeX generated")
    
    # Generate PDF from LaTeX
    pdf_path = service.latex_to_pdf(latex_doc, "sample_assignment")
    if pdf_path:
        print(f"PDF generated: {pdf_path}")
    
    # Generate DOCX
    docx_path = service.generate_docx(content, "Sample Assignment", "John Doe")
    if docx_path:
        print(f"DOCX generated: {docx_path}")

